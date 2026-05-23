"""文档解析服务 — 基于 MinerU (magic-pdf) 的工业级文档解析

架构优先级：
  1. MinerU (magic-pdf CLI) — 用于中文文档、复杂排版、表格提取
  2. PyMuPDF (fitz) — 降级方案，处理普通 PDF
  3. python-docx — DOCX 降级方案

配置：~/magic-pdf.json (含 OpenXLab token)
"""

import json
import logging
import os
import subprocess
import tempfile
import time
from pathlib import Path
from typing import Optional

from app.config import settings

logger = logging.getLogger(__name__)

# 输出目录
PARSER_OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(__file__))),
    "data", "parser_output"
)
os.makedirs(PARSER_OUTPUT_DIR, exist_ok=True)


# ══════════════════════════════════════════════
# MinerU 解析
# ══════════════════════════════════════════════

def _magic_pdf_parse(file_path: str) -> Optional[str]:
    """使用 MinerU (magic-pdf CLI) 解析文档

    特点：自动检测文本/扫描 PDF，支持表格和排版还原
    需要 OpenXLab token 下载模型（已在 ~/magic-pdf.json 中配置）
    """
    try:
        safe_name = os.path.splitext(os.path.basename(file_path))[0]
        safe_name = safe_name.replace(" ", "_").replace("(", "").replace(")", "")
        output_dir = os.path.join(PARSER_OUTPUT_DIR, safe_name)

        # 调用 magic-pdf CLI（auto 模式自动选择最佳方案）
        cmd = [
            "magic-pdf",
            "-p", file_path,
            "-o", output_dir,
            "-m", "auto",
        ]
        logger.info(f"🚀 MinerU 开始解析: {os.path.basename(file_path)}")

        t0 = time.time()
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=300
        )
        elapsed = time.time() - t0

        if result.returncode != 0:
            stderr = result.stderr[:500]
            logger.warning(f"MinerU 解析失败 (exit={result.returncode}, {elapsed:.1f}s): {stderr}")
            return None

        # 递归查找输出 Markdown/Text 文件
        content = _find_output(output_dir, safe_name)
        if content:
            logger.info(f"✅ MinerU 解析成功: {len(content)} 字符 ({elapsed:.1f}s)")
            return content

        logger.warning(f"MinerU 未找到输出文件: {output_dir}")
        return None

    except subprocess.TimeoutExpired:
        logger.error(f"MinerU 超时 (300s): {file_path}")
        return None
    except FileNotFoundError:
        logger.error("magic-pdf CLI 未安装或不在 PATH 中")
        return None
    except Exception as e:
        logger.error(f"MinerU 异常: {e}")
        return None


def _find_output(output_dir: str, base_name: str) -> Optional[str]:
    """在 MinerU 输出目录中递归查找 Markdown/Text 文件"""
    if not os.path.isdir(output_dir):
        return None

    candidates = []

    # 递归查找所有 .md 和 .txt
    for root, dirs, files in os.walk(output_dir):
        for f in files:
            if f.endswith((".md", ".txt")):
                candidates.append(os.path.join(root, f))

    # 优先取最大的 .md 文件
    md_files = [f for f in candidates if f.endswith(".md")]
    if md_files:
        md_files.sort(key=lambda f: os.path.getsize(f), reverse=True)
        with open(md_files[0], "r", encoding="utf-8", errors="ignore") as fh:
            content = fh.read()
        if content.strip():
            return content

    # 降级到 .txt
    txt_files = [f for f in candidates if f.endswith(".txt")]
    if txt_files:
        txt_files.sort(key=lambda f: os.path.getsize(f), reverse=True)
        with open(txt_files[0], "r", encoding="utf-8", errors="ignore") as fh:
            content = fh.read()
        if content.strip():
            return content

    return None


# ══════════════════════════════════════════════
# PyMuPDF 降级方案
# ══════════════════════════════════════════════

def _fitz_parse(file_path: str) -> Optional[str]:
    """PyMuPDF 解析 PDF — 高效处理中文、表格、复杂排版"""
    try:
        import fitz
        doc = fitz.open(file_path)
        pages = []
        for i, page in enumerate(doc):
            # 获取可排序的文本块（保留阅读顺序）
            blocks = page.get_text("blocks")
            page_text = []
            for b in blocks:
                # block format: (x0, y0, x1, y1, text, block_no, block_type)
                if b[6] == 0:  # text block
                    text = b[4].strip()
                    if text:
                        page_text.append(text)
                elif b[6] == 1:  # image block
                    page_text.append("[图片]")

            if page_text:
                pages.append(f"--- 第 {i+1} 页 ---\n" + "\n".join(page_text))

        doc.close()
        content = "\n\n".join(pages) if pages else ""
        return content if content.strip() else None

    except Exception as e:
        logger.warning(f"PyMuPDF 解析失败: {e}")
        return None


# ══════════════════════════════════════════════
# 统一入口
# ══════════════════════════════════════════════

def extract_text(file_path: str, file_type: str) -> str:
    """提取文档文本 — MinerU > PyMuPDF > python-docx

    对 PDF 优先使用 MinerU（支持扫描件、复杂排版），
    失败后降级到 PyMuPDF（支持中文和复杂布局）。
    """
    from app.services.storage import UPLOAD_DIR

    # 构建完整路径
    if not file_path.startswith("/"):
        full_path = os.path.join(UPLOAD_DIR, file_path)
    else:
        full_path = file_path

    if not os.path.exists(full_path):
        logger.error(f"📄 文件不存在: {full_path}")
        return ""

    file_size = os.path.getsize(full_path)
    logger.info(f"📄 开始解析: {os.path.basename(full_path)} ({file_size / 1024:.1f}KB)")

    extracted = ""

    if file_type == "pdf":
        # 第一优先：MinerU
        content = _magic_pdf_parse(full_path)
        if content:
            extracted = content
            logger.info(f"✅ 使用 MinerU 解析成功")
        else:
            # 降级：PyMuPDF
            logger.info("⏬ 降级到 PyMuPDF")
            content = _fitz_parse(full_path)
            if content:
                extracted = content
                logger.info(f"✅ 使用 PyMuPDF 解析成功")

    elif file_type in ("docx", "doc"):
        # 第一优先：MinerU
        content = _magic_pdf_parse(full_path)
        if content:
            extracted = content
            logger.info(f"✅ 使用 MinerU 解析成功")
        else:
            # 降级：python-docx
            logger.info("⏬ 降级到 python-docx")
            try:
                from docx import Document
                doc = Document(full_path)
                texts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        texts.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        texts.append(" | ".join(cell.text.strip() for cell in row.cells))
                extracted = "\n".join(texts)
                if extracted.strip():
                    logger.info(f"✅ 使用 python-docx 解析成功")
            except Exception as e:
                logger.warning(f"python-docx 降级失败: {e}")

    else:
        try:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                extracted = f.read()
        except Exception as e:
            logger.error(f"文本读取失败: {e}")

    if not extracted or not extracted.strip():
        logger.warning(f"⚠️ 解析结果为空: {os.path.basename(full_path)}")
        return ""

    # 清理多余空行并统计
    lines = [l for l in extracted.split("\n") if l.strip()]
    result = "\n".join(lines)
    logger.info(f"✅ 解析完成: {len(result)} 字符, {len(lines)} 行")
    return result
